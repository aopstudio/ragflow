#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#
import re
from copy import deepcopy
from io import BytesIO
from timeit import default_timer as timer
from nltk import word_tokenize
from openpyxl import load_workbook
from rag.nlp import is_english, random_choices, find_codec, qbullets_category, add_positions, has_qbullet
from rag.nlp import rag_tokenizer, tokenize_table
from rag.settings import cron_logger
from docx import Document
from deepdoc.parser import PdfParser, ExcelParser, DocxParser
class Excel(ExcelParser):
    def __call__(self, fnm, binary=None, callback=None):
        if not binary:
            wb = load_workbook(fnm)
        else:
            wb = load_workbook(BytesIO(binary))
        total = 0
        for sheetname in wb.sheetnames:
            total += len(list(wb[sheetname].rows))

        res, fails = [], []
        for sheetname in wb.sheetnames:
            ws = wb[sheetname]
            rows = list(ws.rows)
            for i, r in enumerate(rows):
                q, a = "", ""
                for cell in r:
                    if not cell.value:
                        continue
                    if not q:
                        q = str(cell.value)
                    elif not a:
                        a = str(cell.value)
                    else:
                        break
                if q and a:
                    res.append((q, a))
                else:
                    fails.append(str(i + 1))
                if len(res) % 999 == 0:
                    callback(len(res) *
                             0.6 /
                             total, ("Extract Q&A: {}".format(len(res)) +
                                     (f"{len(fails)} failure, line: %s..." %
                                      (",".join(fails[:3])) if fails else "")))

        callback(0.6, ("Extract Q&A: {}. ".format(len(res)) + (
            f"{len(fails)} failure, line: %s..." % (",".join(fails[:3])) if fails else "")))
        self.is_english = is_english(
            [rmPrefix(q) for q, _ in random_choices(res, k=30) if len(q) > 1])
        return res

class Pdf(PdfParser):
    def __call__(self, filename, binary=None, from_page=0,
                 to_page=100000, zoomin=3, callback=None):
        start = timer()
        callback(msg="OCR is running...")
        self.__images__(
            filename if not binary else binary,
            zoomin,
            from_page,
            to_page,
            callback
        )
        callback(msg="OCR finished")
        cron_logger.info("OCR({}~{}): {}".format(from_page, to_page, timer() - start))
        start = timer()
        self._layouts_rec(zoomin, drop=False)
        callback(0.63, "Layout analysis finished.")
        self._table_transformer_job(zoomin)
        callback(0.65, "Table analysis finished.")
        self._text_merge()
        callback(0.67, "Text merging finished")
        tbls = self._extract_table_figure(True, zoomin, True, True)
        #self._naive_vertical_merge()
        # self._concat_downward()
        #self._filter_forpages()
        cron_logger.info("layouts: {}".format(timer() - start))
        sections = [b["text"] for b in self.boxes]
        bull_x0_list = []
        q_bull, reg = qbullets_category(sections)
        if q_bull == -1:
            raise ValueError("Unable to recognize Q&A structure.")
        qai_list = []
        last_q, last_a, last_tag = '', '', ''
        last_index = -1
        last_box = {'text':''}
        last_bull = None
        for box in self.boxes:
            section, line_tag = box['text'], self._line_tag(box, zoomin)
            has_bull, index = has_qbullet(reg, box, last_box, last_index, last_bull, bull_x0_list)
            last_box, last_index, last_bull = box, index, has_bull
            if not has_bull:  # No question bullet
                if not last_q:
                    continue
                else:
                    last_a = f'{last_a}{section}'
                    last_tag = f'{last_tag}{line_tag}'
            else:
                if last_q:
                    qai_list.append((last_q, last_a, *self.crop(last_tag, need_position=True)))
                    last_q, last_a, last_tag = '', '', ''
                last_q = has_bull.group()
                _, end = has_bull.span()
                last_a = section[end:]
                last_tag = line_tag
        if last_q:
            qai_list.append((last_q, last_a, *self.crop(last_tag, need_position=True)))
        return qai_list, tbls
    
class Docx(DocxParser):
    def __init__(self):
        pass

    def __clean(self, line):
        line = re.sub(r"\u3000", " ", line).strip()
        return line
    def __call__(self, filename, binary=None, from_page=0, to_page=100000, callback=None):
        self.doc = Document(
            filename) if not binary else Document(BytesIO(binary))
        pn = 0
        lines = []
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            if from_page <= pn < to_page and p.text.strip():
                lines.append(self.__clean(p.text))
            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        qa_list = []
        last_question, last_anwser = '', ''
        for line in lines:
            last_pos = 0
            for match in re.finditer(r'(问：|问题：|Q:|Question:)(\w|\W)+(\?|？)',line):  # line以问题开头
                if last_question:
                    last_anwser = f'{last_anwser}{line[last_pos:match.span()[0]]}'
                    qa_list.append((last_question,last_anwser))
                    last_anwser = ''
                last_question = match.group()
                last_pos = match.span()[1]
            last_anwser = f'{last_anwser}{line[last_pos:]}'
        qa_list.append((last_question,last_anwser))
        tbls = []
        for tb in self.doc.tables:
            html= "<table>"
            for r in tb.rows:
                html += "<tr>"
                i = 0
                while i < len(r.cells):
                    span = 1
                    c = r.cells[i]
                    for j in range(i+1, len(r.cells)):
                        if c.text == r.cells[j].text:
                            span += 1
                            i = j
                    i += 1
                    html += f"<td>{c.text}</td>" if span == 1 else f"<td colspan='{span}'>{c.text}</td>"
                html += "</tr>"
            html += "</table>"
            tbls.append(((None, html), ""))
        return qa_list, tbls
    
def rmPrefix(txt):
    return re.sub(
        r"^(问题|答案|回答|user|assistant|Q|A|Question|Answer|问|答)[\t:： ]+", "", txt.strip(), flags=re.IGNORECASE)


def beAdocPdf(d, q, a, eng, image, poss):
    qprefix = "Question: " if eng else "问题："
    aprefix = "Answer: " if eng else "回答："
    d["content_with_weight"] = "\t".join(
        [qprefix + rmPrefix(q), aprefix + rmPrefix(a)])
    d["content_ltks"] = rag_tokenizer.tokenize(q)
    d["content_sm_ltks"] = rag_tokenizer.fine_grained_tokenize(d["content_ltks"])
    d["image"] = image
    add_positions(d, poss)
    return d

def beAdoc(d, q, a, eng):
    qprefix = "Question: " if eng else "问题："
    aprefix = "Answer: " if eng else "回答："
    d["content_with_weight"] = "\t".join(
        [qprefix + rmPrefix(q), aprefix + rmPrefix(a)])
    d["content_ltks"] = rag_tokenizer.tokenize(q)
    d["content_sm_ltks"] = rag_tokenizer.fine_grained_tokenize(d["content_ltks"])
    return d


def chunk(filename, binary=None, lang="Chinese", callback=None, **kwargs):
    """
        Excel and csv(txt) format files are supported.
        If the file is in excel format, there should be 2 column question and answer without header.
        And question column is ahead of answer column.
        And it's O.K if it has multiple sheets as long as the columns are rightly composed.

        If it's in csv format, it should be UTF-8 encoded. Use TAB as delimiter to separate question and answer.

        All the deformed lines will be ignored.
        Every pair of Q&A will be treated as a chunk.
    """
    eng = lang.lower() == "english"
    res = []
    doc = {
        "docnm_kwd": filename,
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))
    }
    if re.search(r"\.xlsx?$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        excel_parser = Excel()
        for q, a in excel_parser(filename, binary, callback):
            res.append(beAdoc(deepcopy(doc), q, a, eng))
        return res
    elif re.search(r"\.(txt|csv)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        txt = ""
        if binary:
            encoding = find_codec(binary)
            txt = binary.decode(encoding, errors="ignore")
        else:
            with open(filename, "r") as f:
                while True:
                    l = f.readline()
                    if not l:
                        break
                    txt += l
        lines = txt.split("\n")
        comma, tab = 0, 0
        for l in lines:
            if len(l.split(",")) == 2: comma += 1
            if len(l.split("\t")) == 2: tab += 1
        delimiter = "\t" if tab >= comma else ","

        fails = []
        question, answer = "", ""
        i = 0
        while i < len(lines):
            arr = lines[i].split(delimiter)
            if len(arr) != 2:
                if question: answer += "\n" + lines[i]
                else:
                    fails.append(str(i+1))
            elif len(arr) == 2:
                if question and answer: res.append(beAdoc(deepcopy(doc), question, answer, eng))
                question, answer = arr
            i += 1
            if len(res) % 999 == 0:
                callback(len(res) * 0.6 / len(lines), ("Extract Q&A: {}".format(len(res)) + (
                    f"{len(fails)} failure, line: %s..." % (",".join(fails[:3])) if fails else "")))

        if question: res.append(beAdoc(deepcopy(doc), question, answer, eng))

        callback(0.6, ("Extract Q&A: {}".format(len(res)) + (
            f"{len(fails)} failure, line: %s..." % (",".join(fails[:3])) if fails else "")))

        return res
    elif re.search(r"\.pdf$", filename, re.IGNORECASE):
        pdf_parser = Pdf()
        count = 0
        qai_list, tbls = pdf_parser(filename if not binary else binary,
                                    from_page=0, to_page=10000, callback=callback)
        
        res = tokenize_table(tbls, doc, eng)

        for q, a, image, poss in qai_list:
            count += 1
            res.append(beAdocPdf(deepcopy(doc), q, a, eng, image, poss))
        return res
    elif re.search(r"\.docx$", filename, re.IGNORECASE):
        docx_parser = Docx()
        qa_list, tbls = docx_parser(filename, binary,
                                    from_page=0, to_page=10000, callback=callback)
        res = tokenize_table(tbls, doc, eng)
        for q, a in qa_list:
            res.append(beAdoc(deepcopy(doc), q, a, eng))
        return res


    raise NotImplementedError(
        "Excel and csv(txt) format files are supported.")


if __name__ == "__main__":
    import sys

    def dummy(prog=None, msg=""):
        pass
    import json
    from icecream import ic
    ic(chunk(sys.argv[1], from_page=0, to_page=10, callback=dummy))
